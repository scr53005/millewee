"""
Millewee QR Code Sticker Generator

Generates print-ready QR code stickers for restaurant tables.
Each sticker contains:
  - Millewee logo (top)
  - "Table: X" label
  - QR code with Innopay logo overlay
  - "Scan to order!" speech bubble
  - URL in small text (bottom)

Uses segno for QR generation, Pillow for composition.
Outputs individual PNGs + a DOCX with 2x2 grid layout.

Usage:
  python generateqrs.py           # all tables
  python generateqrs.py --t       # test mode (last 8 tables)

Prerequisites:
  pip install segno Pillow python-docx

Input files (auto-detected from script directory):
  - {prefix}uri.txt        → base URL
  - {prefix}tables.csv     → table numbers (one per line)
  - logo_millewee_transp.png
  - innologo-71x47.png
"""

import segno
from PIL import Image, ImageDraw, ImageFont, ImageFilter
import os
import sys
import re
import io

# ─── Sticker layout constants (pixels at 300 DPI) ───

# Target sticker: ~5cm x 5.25cm at 300 DPI
STICKER_W = 591   # ~5cm
STICKER_H = 620   # ~5.25cm
DPI = 300

# Colors (Millewee brand)
BG_COLOR = (250, 246, 240)         # warm cream #faf6f0
TEXT_COLOR = (42, 31, 23)          # dark leather #2a1f17
AMBER = (212, 162, 78)            # primary amber #d4a24e
DARK_BROWN = (26, 19, 16)        # #1a1310
BUBBLE_BG = (212, 162, 78)       # amber bubble
BUBBLE_TEXT = (26, 19, 16)        # dark text on amber
URL_COLOR = (42, 31, 23)         # dark leather (same as text)
BORDER_COLOR = (212, 162, 78)    # amber border

# Layout positions
LOGO_Y = 15
LOGO_MAX_H = 90
TABLE_TEXT_Y = 115
QR_SIZE = 340
QR_X = (STICKER_W - QR_SIZE) // 2
QR_Y = 220                        # pushed down for more space after rule
URL_Y = QR_Y + QR_SIZE + 10
BUBBLE_OFFSET_X = -15   # relative to QR left edge
BUBBLE_OFFSET_Y = -62   # fully above QR top edge (+2-3mm higher)


def load_font(name, size, index=0):
    """Try to load a TrueType font, fall back to default."""
    for font_name in [name, "arial.ttf", "Arial.ttf"]:
        try:
            return ImageFont.truetype(font_name, size, index=index)
        except (IOError, OSError):
            continue
    return ImageFont.load_default()


def draw_speech_bubble(img, text, x, y, font, padding=10):
    """Draw a playful comic-style speech bubble with slight tilt and a tail pointing down."""
    # Render bubble on a separate RGBA canvas, then rotate and paste
    temp_w, temp_h = 300, 80
    bubble_canvas = Image.new("RGBA", (temp_w, temp_h), (0, 0, 0, 0))
    bdraw = ImageDraw.Draw(bubble_canvas)

    bbox = bdraw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]

    bx1 = 10
    by1 = 5
    bx2 = bx1 + tw + padding * 2
    by2 = by1 + th + padding * 2
    radius = 14

    # Bubble body (slightly puffy with thicker outline)
    bdraw.rounded_rectangle([bx1, by1, bx2, by2], radius=radius, fill=BUBBLE_BG, outline=DARK_BROWN, width=2)

    # Tail pointing down toward QR
    tail_cx = bx1 + (bx2 - bx1) // 2
    tail_y = by2
    bdraw.polygon([
        (tail_cx - 8, tail_y - 1),
        (tail_cx + 4, tail_y + 14),
        (tail_cx + 14, tail_y - 1),
    ], fill=BUBBLE_BG)
    # Tail outline
    bdraw.line([(tail_cx - 8, tail_y), (tail_cx + 4, tail_y + 14)], fill=DARK_BROWN, width=2)
    bdraw.line([(tail_cx + 4, tail_y + 14), (tail_cx + 14, tail_y)], fill=DARK_BROWN, width=2)

    # Text (bold feel via slight offset double-draw)
    bdraw.text((bx1 + padding, by1 + padding - 2), text, font=font, fill=BUBBLE_TEXT)

    # Rotate slightly for playfulness (-4 degrees)
    bubble_canvas = bubble_canvas.rotate(4, expand=True, resample=Image.BICUBIC)

    # Paste onto main image
    img.paste(bubble_canvas, (x, y), bubble_canvas)


def generate_qr_with_logo(url, qr_size, logo_path):
    """Generate a QR code using segno with the Innopay logo in the center."""
    qr = segno.make(url, error="H")  # High error correction for logo overlay

    # Render to PNG buffer
    buf = io.BytesIO()
    qr.save(buf, kind="png", scale=10, border=0, dark="#2a1f17", light="#ffffff")
    buf.seek(0)
    qr_img = Image.open(buf).convert("RGBA")
    qr_img = qr_img.resize((qr_size, qr_size), Image.LANCZOS)

    # Overlay the Innopay logo in the center
    if os.path.exists(logo_path):
        logo = Image.open(logo_path).convert("RGBA")
        # Logo should be ~20% of QR size
        logo_max = int(qr_size * 0.22)
        logo_ratio = min(logo_max / logo.width, logo_max / logo.height)
        logo_w = int(logo.width * logo_ratio)
        logo_h = int(logo.height * logo_ratio)
        logo = logo.resize((logo_w, logo_h), Image.LANCZOS)

        # White background pad behind logo for readability
        pad = 8
        white_bg = Image.new("RGBA", (logo_w + pad * 2, logo_h + pad * 2), (255, 255, 255, 255))
        lx = (qr_size - white_bg.width) // 2
        ly = (qr_size - white_bg.height) // 2
        qr_img.paste(white_bg, (lx, ly))
        qr_img.paste(logo, (lx + pad, ly + pad), logo)

    return qr_img.convert("RGB")


def create_sticker(table_num, base_uri, millewee_logo_path, inno_logo_path, fonts):
    """Create a single sticker image for a table."""
    font_table, font_url, font_bubble = fonts
    full_url = f"{base_uri}{table_num}"

    img = Image.new("RGB", (STICKER_W, STICKER_H), BG_COLOR)
    draw = ImageDraw.Draw(img)

    # Amber border (rounded rectangle)
    draw.rounded_rectangle(
        [4, 4, STICKER_W - 5, STICKER_H - 5],
        radius=20, outline=BORDER_COLOR, width=3
    )

    # ─── Millewee logo (centered, top) ───
    if os.path.exists(millewee_logo_path):
        logo = Image.open(millewee_logo_path).convert("RGBA")
        ratio = min(LOGO_MAX_H / logo.height, (STICKER_W - 40) / logo.width)
        logo_w = int(logo.width * ratio)
        logo_h = int(logo.height * ratio)
        logo = logo.resize((logo_w, logo_h), Image.LANCZOS)

        # Paste with alpha onto a temp RGBA canvas to handle transparency
        logo_x = (STICKER_W - logo_w) // 2
        temp = Image.new("RGBA", img.size, (0, 0, 0, 0))
        temp.paste(logo, (logo_x, LOGO_Y), logo)
        img = Image.alpha_composite(img.convert("RGBA"), temp).convert("RGB")
        draw = ImageDraw.Draw(img)

    # ─── "Table: X" text (centered) ───
    table_text = f"Table:  {table_num}"
    bbox = draw.textbbox((0, 0), table_text, font=font_table)
    tw = bbox[2] - bbox[0]
    table_text_x = (STICKER_W - tw) // 2
    draw.text((table_text_x, TABLE_TEXT_Y), table_text, font=font_table, fill=TEXT_COLOR)

    # ─── Thin amber line separator under "Table:" ───
    line_y = TABLE_TEXT_Y + (bbox[3] - bbox[1]) + 8
    line_margin = 200
    draw.line([(line_margin, line_y), (STICKER_W - line_margin, line_y)], fill=AMBER, width=2)

    # ─── QR code with Innopay logo ───
    qr_img = generate_qr_with_logo(full_url, QR_SIZE, inno_logo_path)
    img.paste(qr_img, (QR_X, QR_Y))

    # ─── Speech bubble "Scan to order!" (above QR, left side) ───
    bubble_x = QR_X + BUBBLE_OFFSET_X
    bubble_y = QR_Y + BUBBLE_OFFSET_Y
    draw_speech_bubble(img, "Scan to order!", bubble_x, bubble_y, font_bubble)
    draw = ImageDraw.Draw(img)  # refresh draw after paste

    # ─── URL text (small, centered below QR) ───
    url_display = f"https://millewee.innopay.lu/?table={table_num}"
    bbox = draw.textbbox((0, 0), url_display, font=font_url)
    uw = bbox[2] - bbox[0]
    draw.text(((STICKER_W - uw) // 2, URL_Y), url_display, font=font_url, fill=URL_COLOR)

    return img


def generate_all(uri_file, tables_file, script_dir, test_mode=False):
    """Main generation pipeline."""

    # Read inputs
    with open(uri_file, "r", encoding="utf-8-sig") as f:
        base_uri = f.read().strip()

    with open(tables_file, "r", encoding="utf-8-sig") as f:
        table_numbers = [line.strip() for line in f if line.strip()]

    if test_mode:
        table_numbers = table_numbers[-4:]
        print(f"Test mode: {len(table_numbers)} tables only.")

    # Paths
    millewee_logo = os.path.join(script_dir, "logo_millewee_transp.png")
    inno_logo = os.path.join(script_dir, "innologo-71x47.png")
    output_dir = os.path.join(script_dir, "output_qrs")
    os.makedirs(output_dir, exist_ok=True)

    # Fonts
    SITKA = "C:/Windows/Fonts/Sitka.ttc"
    font_table = load_font(SITKA, 38)
    font_url = load_font(SITKA, 19)
    font_bubble = load_font(SITKA, 19)
    fonts = (font_table, font_url, font_bubble)

    print(f"Generating {len(table_numbers)} stickers...\n")

    sticker_paths = []
    for table_num in table_numbers:
        sticker = create_sticker(table_num, base_uri, millewee_logo, inno_logo, fonts)
        out_path = os.path.join(output_dir, f"table_{table_num.zfill(2)}.png")
        sticker.save(out_path, "PNG", dpi=(DPI, DPI))
        sticker_paths.append(out_path)
        print(f"  OK Table {table_num}")

    # ─── Compile into DOCX (3x2 landscape grid = 6 per page) ───
    # Sticker aspect ratio: 945 x 709 px = ~8cm x 6cm
    # A4 landscape: 29.7cm x 21cm, with tight margins → ~27cm x 19cm usable
    # 3 cols × 8.8cm = 26.4cm, 2 rows × 9.3cm = 18.6cm → fits!
    try:
        from docx import Document
        from docx.shared import Cm, Pt, Twips
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        COLS = 5
        ROWS = 3
        PER_PAGE = COLS * ROWS

        # Sticker image width in the doc (leave a bit of cell padding)
        IMG_W_CM = 5.2

        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

        def _remove_borders(table):
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl.insert(0, tblPr)
            borders = OxmlElement("w:tblBorders")
            for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                elem = OxmlElement(f"w:{name}")
                elem.set(qn("w:val"), "nil")
                elem.set(qn("w:sz"), "0")
                elem.set(qn("w:space"), "0")
                borders.append(elem)
            tblPr.append(borders)

        def _clear_cell_margins(cell):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for name in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{name}")
                node.set(qn("w:w"), "20")   # ~0.35mm padding
                node.set(qn("w:type"), "dxa")
                tcMar.append(node)
            tcPr.append(tcMar)

        for i in range(0, len(sticker_paths), PER_PAGE):
            table = doc.add_table(rows=ROWS, cols=COLS)
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_borders(table)

            for row_idx in range(ROWS):
                for col_idx in range(COLS):
                    img_idx = i + row_idx * COLS + col_idx
                    cell = table.cell(row_idx, col_idx)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    _clear_cell_margins(cell)

                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)

                    if img_idx < len(sticker_paths):
                        run = para.add_run()
                        run.add_picture(sticker_paths[img_idx], width=Cm(IMG_W_CM))

            # Page break between groups (except last)
            if i + PER_PAGE < len(sticker_paths):
                doc.add_page_break()

        suffix = "_test" if test_mode else ""
        docx_path = os.path.join(output_dir, f"Millewee_QR_Stickers{suffix}.docx")

        while True:
            try:
                doc.save(docx_path)
                print(f"\nOK DOCX saved: {docx_path}")
                break
            except PermissionError:
                print(f"\nCannot save '{docx_path}' — close it in Word first.")
                input("Press Enter to retry...")

    except ImportError:
        print("\npython-docx not installed — skipping DOCX generation.")
        print("  pip install python-docx")

    print(f"\nDone! {len(sticker_paths)} stickers in {output_dir}/")


# ─── Auto-discover and run ───

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    uri_pattern = re.compile(r"^([a-z0-9]+)uri\.txt$", re.IGNORECASE)

    prefix = None
    uri_file = None
    for filename in os.listdir(script_dir):
        match = uri_pattern.match(filename)
        if match:
            prefix = match.group(1)
            uri_file = os.path.join(script_dir, filename)
            break

    if uri_file is None:
        print("Error: no <prefix>uri.txt found in script directory.")
        sys.exit(1)

    print(f"Detected prefix: '{prefix}'")
    tables_file = os.path.join(script_dir, f"{prefix}tables.csv")

    test_mode = "--t" in sys.argv
    if test_mode:
        print("Test mode enabled (--t)")

    generate_all(uri_file, tables_file, script_dir, test_mode=test_mode)
